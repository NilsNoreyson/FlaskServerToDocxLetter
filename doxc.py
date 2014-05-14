# -*- coding: utf-8 -*-
"""
Created on Tue May  6 11:00:11 2014

@author: peterb
"""
#pip install python-docx
#from https://python-docx.readthedocs.org/en/latest/user/install.html
from docx import Document
from docx.shared import Inches


from IMAP import *

    
    

IMAP=openIMAPConnection(SERVER,PORT,USER,PASSWORD,FOLDER)
ids=getIDS(IMAP)

nameEmail=[]
for i in ids:
    vCard=getIdVcards(i,IMAP)
    print vCard
    di=parseVCard(vCard)
    try:
        email=di['email']['text']
        name=di['fn']['text']
        nameEmail.append("%s---%s" %(name,email))
    except:
        pass

closeIMAP(IMAP)





document = Document()

document.add_heading('Document Title', 0)

p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='IntenseQuote')

document.add_paragraph(
    'first item in unordered list', style='ListBullet'
)

for s in nameEmail:
    document.add_paragraph(
        s, style='ListNumber'
    )




document.add_page_break()


document.save('/home/peterb/demo.docx')